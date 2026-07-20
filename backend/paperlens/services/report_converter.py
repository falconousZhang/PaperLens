from __future__ import annotations

import io
import re
import zipfile
from xml.etree import ElementTree

from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
)
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


_FIXED_CREATOR = "PaperLens"
_FIXED_PRODUCER = "PaperLens Report Generator"
_PDF_FONT_NAME = "STSong-Light"
_NAVY = colors.HexColor("#1F2A44")
_BLUE = colors.HexColor("#2F5D8C")
_LIGHT_BLUE = colors.HexColor("#EAF0F6")
_LIGHT_GRAY = colors.HexColor("#F4F6F8")
_TEXT = colors.HexColor("#243142")

pdfmetrics.registerFont(UnicodeCIDFont(_PDF_FONT_NAME))
pdfmetrics.registerFontFamily(
    _PDF_FONT_NAME,
    normal=_PDF_FONT_NAME,
    bold=_PDF_FONT_NAME,
    italic=_PDF_FONT_NAME,
    boldItalic=_PDF_FONT_NAME,
)


def markdown_to_pdf(md_bytes: bytes) -> bytes:
    md_text = md_bytes.decode("utf-8")
    lines = md_text.split("\n")
    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=LETTER,
        leftMargin=inch,
        rightMargin=inch,
        topMargin=0.82 * inch,
        bottomMargin=0.78 * inch,
        title="PaperLens Learning Report",
        author=_FIXED_CREATOR,
        subject="Paper Learning Report",
        creator=_FIXED_CREATOR,
        invariant=1,
    )
    styles = getSampleStyleSheet()
    h1_style = ParagraphStyle(
        "H1Custom",
        parent=styles["Heading1"],
        fontSize=22,
        fontName=_PDF_FONT_NAME,
        textColor=_NAVY,
        leading=28,
        spaceAfter=16,
        spaceBefore=6,
    )
    h2_style = ParagraphStyle(
        "H2Custom",
        parent=styles["Heading2"],
        fontSize=14,
        fontName=_PDF_FONT_NAME,
        textColor=_BLUE,
        leading=18,
        spaceAfter=7,
        spaceBefore=15,
        keepWithNext=True,
    )
    h3_style = ParagraphStyle(
        "H3Custom",
        parent=styles["Heading3"],
        fontSize=11.5,
        fontName=_PDF_FONT_NAME,
        textColor=_NAVY,
        leading=15,
        spaceAfter=5,
        spaceBefore=10,
        keepWithNext=True,
    )
    body_style = ParagraphStyle(
        "BodyCustom",
        parent=styles["Normal"],
        fontSize=9.5,
        fontName=_PDF_FONT_NAME,
        textColor=_TEXT,
        spaceAfter=5,
        leading=14.2,
    )
    bullet_style = ParagraphStyle(
        "BulletCustom",
        parent=styles["Normal"],
        fontSize=9.5,
        fontName=_PDF_FONT_NAME,
        textColor=_TEXT,
        spaceAfter=3,
        leftIndent=14,
        firstLineIndent=-8,
        leading=14.2,
    )

    story = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("### "):
            text = _strip_md_formatting(stripped[4:])
            story.append(Paragraph(_esc_pdf(text), h3_style))
        elif stripped.startswith("## "):
            text = _strip_md_formatting(stripped[3:])
            story.append(Paragraph(_esc_pdf(text), h2_style))
        elif stripped.startswith("# "):
            text = _strip_md_formatting(stripped[2:])
            story.append(Paragraph(_esc_pdf(text), h1_style))
        elif stripped.startswith("---"):
            pass
        elif stripped.startswith("| ") and "---" not in stripped:
            table_lines = [stripped]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith("|"):
                if "---" in lines[j]:
                    j += 1
                    continue
                table_lines.append(lines[j].strip())
                j += 1
            i = j - 1
            story.append(_build_pdf_table(table_lines, body_style))
            story.append(Spacer(1, 4))
        elif stripped.startswith("- "):
            text = _strip_md_formatting(stripped[2:])
            story.append(Paragraph(f"\u2022 {_esc_pdf(text)}", bullet_style))
        elif stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            text = _strip_md_formatting(stripped[1:-1])
            story.append(Paragraph(f"<i>{_esc_pdf(text)}</i>", body_style))
        elif stripped.startswith("**") and stripped.endswith("**"):
            text = _strip_md_formatting(stripped[2:-2])
            story.append(Paragraph(f"<b>{_esc_pdf(text)}</b>", body_style))
        else:
            text = _strip_md_formatting(stripped)
            story.append(Paragraph(_esc_pdf(text), body_style))

        i += 1

    doc.build(story, onFirstPage=_add_page_number, onLaterPages=_add_page_number)
    return buf.getvalue()


def markdown_to_docx(md_bytes: bytes) -> bytes:
    md_text = md_bytes.decode("utf-8")
    lines = md_text.split("\n")
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    _set_docx_style_font(style, "Calibri", "Microsoft YaHei", 10.5, "243142")
    style.paragraph_format.space_after = Pt(5)
    style.paragraph_format.line_spacing = 1.15
    _set_docx_style_font(doc.styles["Heading 1"], "Calibri", "Microsoft YaHei", 22, "1F2A44")
    _set_docx_style_font(doc.styles["Heading 2"], "Calibri", "Microsoft YaHei", 14, "2F5D8C")
    _set_docx_style_font(doc.styles["Heading 3"], "Calibri", "Microsoft YaHei", 11.5, "1F2A44")
    doc.styles["Heading 1"].paragraph_format.space_after = Pt(14)
    doc.styles["Heading 2"].paragraph_format.space_before = Pt(12)
    doc.styles["Heading 2"].paragraph_format.space_after = Pt(6)
    doc.styles["Heading 3"].paragraph_format.space_before = Pt(8)
    doc.styles["Heading 3"].paragraph_format.space_after = Pt(4)
    _configure_docx_header_footer(section)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        if stripped.startswith("### "):
            text = _strip_md_formatting(stripped[4:])
            doc.add_heading(text, level=3)
        elif stripped.startswith("## "):
            text = _strip_md_formatting(stripped[3:])
            doc.add_heading(text, level=2)
        elif stripped.startswith("# "):
            text = _strip_md_formatting(stripped[2:])
            doc.add_heading(text, level=1)
        elif stripped.startswith("---"):
            pass
        elif stripped.startswith("| ") and "---" not in stripped:
            table_lines = [stripped]
            j = i + 1
            while j < len(lines) and lines[j].strip().startswith("|"):
                if "---" in lines[j]:
                    j += 1
                    continue
                table_lines.append(lines[j].strip())
                j += 1
            i = j - 1
            _build_docx_table(doc, table_lines)
        elif stripped.startswith("- "):
            text = _strip_md_formatting(stripped[2:])
            doc.add_paragraph(text, style="List Bullet")
        elif stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            text = _strip_md_formatting(stripped[1:-1])
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.italic = True
        elif stripped.startswith("**") and stripped.endswith("**"):
            text = _strip_md_formatting(stripped[2:-2])
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.bold = True
        else:
            text = _strip_md_formatting(stripped)
            doc.add_paragraph(text)

        i += 1

    core_props = doc.core_properties
    core_props.author = _FIXED_CREATOR
    core_props.title = "PaperLens Learning Report"
    core_props.subject = "Paper Learning Report"
    core_props.creator = _FIXED_CREATOR
    core_props.keywords = None
    core_props.comments = None
    core_props.category = None
    core_props.content_status = None
    core_props.identifier = None
    core_props.language = None
    core_props.last_modified_by = None
    core_props.revision = 1

    buf = io.BytesIO()
    doc.save(buf)
    raw = buf.getvalue()
    return _make_deterministic_docx(raw)


def _esc_pdf(text: str) -> str:
    return text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _strip_md_formatting(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"`(.+?)`", r"\1", text)
    return text


def _set_docx_style_font(style, latin: str, east_asia: str, size: float, color: str) -> None:
    style.font.name = latin
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def _configure_docx_header_footer(section) -> None:
    header = section.header
    header.is_linked_to_previous = False
    header_paragraph = header.paragraphs[0]
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_paragraph.add_run("PaperLens · 论文学习报告")
    header_run.font.name = "Calibri"
    header_run.font.size = Pt(8.5)
    header_run.font.color.rgb = RGBColor.from_string("607086")
    header_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string("607086")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instruction, separate, text, end):
        run._element.append(node)


def _add_page_number(canvas, doc):
    canvas.saveState()
    canvas.setAuthor(_FIXED_CREATOR)
    canvas.setCreator(_FIXED_CREATOR)
    canvas.setProducer(_FIXED_PRODUCER)
    canvas.setSubject("Paper Learning Report")
    canvas.setTitle("PaperLens Learning Report")
    canvas.setStrokeColor(colors.HexColor("#D8DEE8"))
    canvas.setLineWidth(0.4)
    canvas.line(inch, 0.58 * inch, LETTER[0] - inch, 0.58 * inch)
    canvas.setFont(_PDF_FONT_NAME, 9)
    canvas.setFillColor(colors.HexColor("#607086"))
    page_num = canvas.getPageNumber()
    canvas.drawCentredString(LETTER[0] / 2, 0.38 * inch, str(page_num))
    canvas.restoreState()


def _build_pdf_table(table_lines: list[str], style) -> Table:
    rows = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows.append([Paragraph(_esc_pdf(c), style) for c in cells])
    if not rows:
        return Spacer(1, 0)
    col_count = max(len(row) for row in rows)
    for row in rows:
        row.extend([Paragraph("", style)] * (col_count - len(row)))
    t = Table(
        rows,
        colWidths=[6.5 * inch / col_count] * col_count,
        repeatRows=1,
        hAlign="LEFT",
    )
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), _NAVY),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("BACKGROUND", (0, 1), (-1, -1), _LIGHT_GRAY),
        ("GRID", (0, 0), (-1, -1), 0.35, colors.HexColor("#D8DEE8")),
        ("FONTSIZE", (0, 0), (-1, -1), 8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    return t


def _build_docx_table(doc: Document, table_lines: list[str]) -> None:
    rows_data = []
    for line in table_lines:
        cells = [c.strip() for c in line.strip("|").split("|")]
        rows_data.append(cells)
    if not rows_data:
        return
    col_count = len(rows_data[0])
    table = doc.add_table(rows=len(rows_data), cols=col_count)
    table.style = "Light Shading Accent 1"
    for row_idx, row_data in enumerate(rows_data):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < col_count:
                cell = table.cell(row_idx, col_idx)
                cell.text = _strip_md_formatting(cell_text)
                if row_idx == 0:
                    shading = OxmlElement("w:shd")
                    shading.set(qn("w:fill"), "1F2A44")
                    cell._tc.get_or_add_tcPr().append(shading)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(255, 255, 255)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = "Calibri"
                        run.font.size = Pt(8.5)
                        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _make_deterministic_docx(raw: bytes) -> bytes:
    buf_in = io.BytesIO(raw)
    buf_out = io.BytesIO()
    with zipfile.ZipFile(buf_in, "r") as zin:
        with zipfile.ZipFile(buf_out, "w", zipfile.ZIP_DEFLATED) as zout:
            for info in sorted(zin.infolist(), key=lambda x: x.filename):
                data = zin.read(info.filename)
                if info.filename == "docProps/core.xml":
                    data = _fix_core_xml(data)
                elif info.filename.startswith("word/") and info.filename.endswith(".xml"):
                    data = re.sub(rb'\s+w:rsid[A-Za-z]*="[^"]*"', b"", data)
                    data = re.sub(rb"<w:rsids\b[^>]*>.*?</w:rsids>", b"", data, flags=re.DOTALL)
                    data = re.sub(rb"<w:rsid[A-Za-z]*\b[^>]*/>", b"", data)
                new_info = zipfile.ZipInfo(info.filename, date_time=(2026, 1, 1, 0, 0, 0))
                new_info.compress_type = zipfile.ZIP_DEFLATED
                new_info.create_system = 0
                new_info.external_attr = 0
                zout.writestr(new_info, data)
    result = buf_out.getvalue()
    _validate_docx_package(result)
    return result


def _fix_core_xml(data: bytes) -> bytes:
    root = ElementTree.fromstring(data)
    ns = {"cp": "http://schemas.openxmlformats.org/package/2006/metadata/core-properties",
          "dc": "http://purl.org/dc/elements/1.1/",
          "dcterms": "http://purl.org/dc/terms/",
          "dcmitype": "http://purl.org/dc/dcmitype/",
          "xsi": "http://www.w3.org/2001/XMLSchema-instance"}
    for tag in ["cp:lastModifiedBy", "cp:revision"]:
        elem = root.find(tag, ns)
        if elem is not None:
            root.remove(elem)
    for tag in ["dcterms:created", "dcterms:modified"]:
        elem = root.find(tag, ns)
        if elem is not None:
            elem.text = "2026-01-01T00:00:00Z"
    ElementTree.register_namespace("cp", ns["cp"])
    ElementTree.register_namespace("dc", ns["dc"])
    ElementTree.register_namespace("dcterms", ns["dcterms"])
    ElementTree.register_namespace("dcmitype", ns["dcmitype"])
    ElementTree.register_namespace("xsi", ns["xsi"])
    return ElementTree.tostring(root, encoding="unicode", xml_declaration=True).encode("utf-8")


def _validate_docx_package(content: bytes) -> None:
    with zipfile.ZipFile(io.BytesIO(content)) as archive:
        names = archive.namelist()
        lowered = [name.lower() for name in names]
        if any("vbaproject" in name or "oleobject" in name or "/embeddings/" in name for name in lowered):
            raise ValueError("unsafe DOCX package")
        for name in names:
            if not name.endswith(".rels"):
                continue
            root = ElementTree.fromstring(archive.read(name))
            for relationship in root:
                if relationship.attrib.get("TargetMode", "").lower() == "external":
                    raise ValueError("external DOCX relationship")
