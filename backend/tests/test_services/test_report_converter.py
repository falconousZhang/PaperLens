import hashlib
import io
import zipfile
from types import SimpleNamespace

import fitz
import pytest
from docx import Document

from paperlens.services.report_converter import (
    markdown_to_pdf,
    markdown_to_docx,
    _FIXED_CREATOR,
    _FIXED_PRODUCER,
)


def _make_md(**overrides):
    parts = ["# Test Report", "", "## Paper Information", "", "- **Title**: Test Paper", "- **Filename**: test.pdf", ""]
    if overrides.get("with_review", True):
        parts += ["## Review Details", "", "### SOUNDNESS", "", "**Rating**: 4/5", "", "**Summary**: Good methodology", ""]
    if overrides.get("with_metrics", False):
        parts += ["## Metrics Data", "", "| Model | Dataset | Metric | Value | Checkpoint |", "| --- | --- | --- | --- | --- |", "| BERT | SQuAD | Accuracy | 92.5 | Best |", ""]
    if overrides.get("with_experiment", False):
        parts += ["## Experiment Analysis Data", "", "### exp.csv", "", "- **Rows**: 100", "- **Columns**: 5", ""]
    return "\n".join(parts).encode("utf-8")


class TestPDFConverter:
    def test_pdf_starts_with_signature(self):
        pdf = markdown_to_pdf(b"# Hello\n")
        assert pdf[:5] == b"%PDF-"

    def test_pdf_deterministic(self):
        md = _make_md()
        pdf1 = markdown_to_pdf(md)
        pdf2 = markdown_to_pdf(md)
        assert pdf1 == pdf2

    def test_pdf_deterministic_hash(self):
        md = _make_md()
        pdf1 = markdown_to_pdf(md)
        pdf2 = markdown_to_pdf(md)
        assert hashlib.sha256(pdf1).hexdigest() == hashlib.sha256(pdf2).hexdigest()

    def test_pdf_text_extractable(self):
        md = _make_md()
        pdf = markdown_to_pdf(md)
        doc = fitz.open(stream=pdf, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        assert "Test Report" in text
        assert "SOUNDNESS" in text

    def test_pdf_zh_text(self):
        md = "# 论文审阅报告\n\n## 审阅详情\n\n### SOUNDNESS\n\n**摘要**: 好的方法\n".encode("utf-8")
        pdf = markdown_to_pdf(md)
        doc = fitz.open(stream=pdf, filetype="pdf")
        text = doc[0].get_text()
        doc.close()
        assert "论文审阅报告" in text
        assert "审阅详情" in text
        assert "摘要" in text
        assert "好的方法" in text

    def test_pdf_has_pages(self):
        md = _make_md()
        pdf = markdown_to_pdf(md)
        doc = fitz.open(stream=pdf, filetype="pdf")
        assert doc.page_count >= 1
        doc.close()

    def test_pdf_no_javascript(self):
        pdf = markdown_to_pdf(b"# Test\n")
        assert b"/JavaScript" not in pdf
        assert b"/JS" not in pdf

    def test_pdf_no_external_resources(self):
        pdf = markdown_to_pdf(b"# Test\n")
        assert b"/EmbeddedFile" not in pdf

    def test_pdf_fixed_creator(self):
        pdf = markdown_to_pdf(b"# Test\n")
        doc = fitz.open(stream=pdf, filetype="pdf")
        assert doc.metadata["author"] == _FIXED_CREATOR
        assert doc.metadata["creator"] == _FIXED_CREATOR
        assert doc.metadata["producer"] == _FIXED_PRODUCER
        doc.close()

    def test_pdf_fixed_dates(self):
        pdf = markdown_to_pdf(b"# Test\n")
        doc = fitz.open(stream=pdf, filetype="pdf")
        assert doc.metadata["creationDate"] == "D:20000101000000+00'00'"
        assert doc.metadata["modDate"] == "D:20000101000000+00'00'"
        doc.close()

    def test_pdf_with_table(self):
        md = b"# Test\n\n| A | B |\n| --- | --- |\n| 1 | 2 |\n| 3 | 4 |\n"
        pdf = markdown_to_pdf(md)
        doc = fitz.open(stream=pdf, filetype="pdf")
        text = doc[0].get_text()
        doc.close()
        assert "A" in text or "1" in text

    def test_pdf_with_bullet_list(self):
        md = b"# Test\n\n- item 1\n- item 2\n- item 3\n"
        pdf = markdown_to_pdf(md)
        doc = fitz.open(stream=pdf, filetype="pdf")
        text = doc[0].get_text()
        doc.close()
        assert "item 1" in text

    def test_pdf_unicode(self):
        md = "# 深度学习\n\n**摘要**: 好的方法\n".encode("utf-8")
        pdf = markdown_to_pdf(md)
        doc = fitz.open(stream=pdf, filetype="pdf")
        text = "".join(page.get_text() for page in doc)
        doc.close()
        assert "深度学习" in text
        assert "好的方法" in text

    def test_pdf_has_no_unsafe_actions_or_attachments(self):
        pdf = markdown_to_pdf(b"# Test\n")
        for marker in [b"/JavaScript", b"/JS", b"/OpenAction", b"/Launch", b"/EmbeddedFile", b"/Filespec"]:
            assert marker not in pdf

    def test_pdf_long_text_multiple_pages(self):
        lines = ["# Long Report", ""]
        for i in range(200):
            lines.append(f"Paragraph {i}: " + "x" * 80)
            lines.append("")
        md = "\n".join(lines).encode("utf-8")
        pdf = markdown_to_pdf(md)
        doc = fitz.open(stream=pdf, filetype="pdf")
        assert doc.page_count > 1
        doc.close()

    def test_pdf_empty_sections(self):
        md = b"# Test\n\n## Empty Section\n\n## Next Section\n\nContent here\n"
        pdf = markdown_to_pdf(md)
        assert len(pdf) > 100

    def test_pdf_no_secrets(self):
        md = b"# Test\n"
        pdf = markdown_to_pdf(md)
        text = pdf.decode("latin-1", errors="replace")
        forbidden = ["storage_key", "content_hash", "raw_text", "/app/", "api_key", "token"]
        for f in forbidden:
            assert f not in text.lower(), f"Forbidden '{f}' in PDF"


class TestDOCXConverter:
    def test_docx_is_valid_zip(self):
        docx = markdown_to_docx(b"# Hello\n")
        assert docx[:4] == b"PK\x03\x04"

    def test_docx_deterministic(self):
        md = _make_md()
        docx1 = markdown_to_docx(md)
        docx2 = markdown_to_docx(md)
        assert docx1 == docx2

    def test_docx_deterministic_hash(self):
        md = _make_md()
        docx1 = markdown_to_docx(md)
        docx2 = markdown_to_docx(md)
        assert hashlib.sha256(docx1).hexdigest() == hashlib.sha256(docx2).hexdigest()

    def test_docx_can_be_reopened(self):
        md = _make_md()
        docx = markdown_to_docx(md)
        doc = Document(io.BytesIO(docx))
        assert len(doc.paragraphs) > 0

    def test_docx_has_headings(self):
        md = b"# Title\n\n## Section\n\n### Subsection\n"
        docx = markdown_to_docx(md)
        doc = Document(io.BytesIO(docx))
        heading_styles = [p.style.name for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(heading_styles) >= 2

    def test_docx_has_table(self):
        md = b"# Test\n\n| A | B |\n| --- | --- |\n| 1 | 2 |\n"
        docx = markdown_to_docx(md)
        doc = Document(io.BytesIO(docx))
        assert len(doc.tables) == 1

    def test_docx_has_bullet_list(self):
        md = b"# Test\n\n- item 1\n- item 2\n"
        docx = markdown_to_docx(md)
        doc = Document(io.BytesIO(docx))
        bullet_paras = [p for p in doc.paragraphs if p.style.name == "List Bullet"]
        assert len(bullet_paras) >= 2

    def test_docx_no_macros(self):
        docx = markdown_to_docx(b"# Test\n")
        with zipfile.ZipFile(io.BytesIO(docx)) as zf:
            names = zf.namelist()
            assert not any("vba" in n.lower() for n in names)
            assert not any("macro" in n.lower() for n in names)

    def test_docx_no_external_relationships(self):
        docx = markdown_to_docx(b"# Test\n")
        with zipfile.ZipFile(io.BytesIO(docx)) as zf:
            for name in zf.namelist():
                if name.endswith(".rels"):
                    assert b'TargetMode="External"' not in zf.read(name)

    def test_docx_no_ole_or_embedded_objects(self):
        docx = markdown_to_docx(b"# Test\n")
        with zipfile.ZipFile(io.BytesIO(docx)) as zf:
            names = [name.lower() for name in zf.namelist()]
            assert not any("oleobject" in name for name in names)
            assert not any("/embeddings/" in name for name in names)

    def test_docx_has_no_revision_session_ids(self):
        docx = markdown_to_docx(_make_md(with_metrics=True))
        with zipfile.ZipFile(io.BytesIO(docx)) as zf:
            for name in zf.namelist():
                if name.startswith("word/") and name.endswith(".xml"):
                    assert b"w:rsid" not in zf.read(name)

    def test_docx_deterministic_zip_timestamps(self):
        md = _make_md()
        docx = markdown_to_docx(md)
        with zipfile.ZipFile(io.BytesIO(docx)) as zf:
            for info in zf.infolist():
                assert info.date_time == (2026, 1, 1, 0, 0, 0)

    def test_docx_sorted_entries(self):
        md = _make_md()
        docx = markdown_to_docx(md)
        with zipfile.ZipFile(io.BytesIO(docx)) as zf:
            names = zf.namelist()
            assert names == sorted(names)

    def test_docx_unicode(self):
        md = "# 深度学习\n\n**摘要**: 好的方法\n".encode("utf-8")
        docx = markdown_to_docx(md)
        doc = Document(io.BytesIO(docx))
        assert any("深度学习" in p.text for p in doc.paragraphs)

    def test_docx_no_secrets(self):
        md = b"# Test\n"
        docx = markdown_to_docx(md)
        with zipfile.ZipFile(io.BytesIO(docx)) as zf:
            for name in zf.namelist():
                if name.endswith(".xml"):
                    content = zf.read(name).decode("utf-8", errors="replace").lower()
                    forbidden = ["storage_key", "content_hash", "api_key", "token", "raw_text"]
                    for f in forbidden:
                        assert f not in content, f"Forbidden '{f}' in {name}"


class TestCrossFormatConsistency:
    def test_same_md_produces_consistent_hashes(self):
        md = _make_md()
        pdf = markdown_to_pdf(md)
        docx = markdown_to_docx(md)
        pdf_hash = hashlib.sha256(pdf).hexdigest()
        docx_hash = hashlib.sha256(docx).hexdigest()
        assert pdf_hash != docx_hash
        assert len(pdf_hash) == 64
        assert len(docx_hash) == 64

    def test_different_md_produces_different_hashes(self):
        md1 = b"# Report A\n\nContent A\n"
        md2 = b"# Report B\n\nContent B\n"
        pdf1 = markdown_to_pdf(md1)
        pdf2 = markdown_to_pdf(md2)
        assert hashlib.sha256(pdf1).hexdigest() != hashlib.sha256(pdf2).hexdigest()
